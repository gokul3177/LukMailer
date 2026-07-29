"""
backend/parser.py — HR email list parser for .csv and .docx files.

Parses recruiter contacts, performs email validation and receiver existence checks,
and returns structured contacts along with comprehensive statistics.
"""

import csv
import io
import re
from pathlib import Path
import docx  # python-docx

from backend.validators import validate_email_syntax, validate_file_extension
from backend.verifier import verify_email_existence

# Regex to detect emails in raw text / docx files
EMAIL_REGEX_PATTERN = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")

def _domain_to_company(email: str) -> str:
    try:
        domain = email.split("@")[1].split(".")[0]
        return domain.replace("-", " ").title()
    except Exception:
        return "Target Company"

def parse_csv_file(file_content: bytes) -> list[tuple[str | None, str | None, str]]:
    """
    Parses CSV content. Supports headers like 'email', 'company', 'name', 'hr_name'
    or raw columns (company|hr_name|email or first|last|email).
    Returns list of raw tuples (company, hr_name, email).
    """
    text = file_content.decode("utf-8", errors="replace")
    stream = io.StringIO(text)
    
    # Check if pipe separated or comma separated
    sample = text[:1024]
    delimiter = "|" if "|" in sample and "," not in sample else ","
    
    reader = csv.reader(stream, delimiter=delimiter)
    raw_contacts = []
    
    header = None
    for row in reader:
        if not row or not any(row):
            continue
        row_str = [cell.strip() for cell in row]
        
        # Check if first row is header
        if header is None and any(col.lower() in ("email", "e-mail", "company", "hr", "name") for col in row_str):
            header = [col.lower() for col in row_str]
            continue
            
        if header:
            # Column mapping
            email, company, hr_name = None, None, None
            for idx, col_name in enumerate(header):
                if idx >= len(row_str):
                    continue
                val = row_str[idx]
                if "email" in col_name or "e-mail" in col_name:
                    email = val
                elif "company" in col_name or "organization" in col_name:
                    company = val
                elif "name" in col_name or "hr" in col_name or "recruiter" in col_name:
                    hr_name = val
            if email:
                raw_contacts.append((company, hr_name, email))
        else:
            # Fallback based on column counts
            if len(row_str) >= 3:
                if "@" in row_str[2]:
                    # first|last|email or company|name|email
                    if "@" not in row_str[0]:
                        c_or_f = row_str[0]
                        last = row_str[1]
                        hr = f"{c_or_f} {last}".strip() if last else c_or_f
                        raw_contacts.append((None, hr, row_str[2]))
                elif "@" in row_str[0]:
                    raw_contacts.append((row_str[1], row_str[2], row_str[0]))
            elif len(row_str) == 2:
                if "@" in row_str[1]:
                    raw_contacts.append((None, row_str[0], row_str[1]))
                elif "@" in row_str[0]:
                    raw_contacts.append((None, row_str[1], row_str[0]))
            elif len(row_str) == 1 and "@" in row_str[0]:
                raw_contacts.append((None, None, row_str[0]))
                
    return raw_contacts

def parse_docx_file(file_content: bytes) -> list[tuple[str | None, str | None, str]]:
    """
    Parses DOCX file content. Extracts emails from paragraphs and tables.
    Returns raw tuples (company, hr_name, email).
    """
    stream = io.BytesIO(file_content)
    doc = docx.Document(stream)
    
    raw_contacts = []
    
    # Process paragraphs
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        found_emails = EMAIL_REGEX_PATTERN.findall(text)
        for email in found_emails:
            raw_contacts.append((None, None, email))
            
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            row_text = " ".join(cells)
            found_emails = EMAIL_REGEX_PATTERN.findall(row_text)
            for email in found_emails:
                raw_contacts.append((None, None, email))
                
    return raw_contacts

def parse_hr_email_list(
    filename: str,
    file_content: bytes,
    verify_existence: bool = True
) -> tuple[list[dict], dict]:
    """
    Main parser function for uploaded HR email lists (.csv or .docx).
    
    Returns:
        (valid_contacts_list, stats_dict)
    """
    valid_ext, err_msg = validate_file_extension(filename, {".csv", ".docx"})
    if not valid_ext:
        raise ValueError(err_msg)

    ext = Path(filename).suffix.lower()
    if ext == ".csv":
        raw_entries = parse_csv_file(file_content)
    elif ext == ".docx":
        raw_entries = parse_docx_file(file_content)
    else:
        raise ValueError("Unsupported file format.")

    total_emails_found = len(raw_entries)
    processed_contacts = []
    seen_emails = set()

    valid_count = 0
    invalid_count = 0
    skipped_duplicates = 0

    for company, hr_name, email in raw_entries:
        if not email or not isinstance(email, str):
            invalid_count += 1
            continue
            
        clean_email = email.strip()
        
        # Check syntax
        if not validate_email_syntax(clean_email):
            invalid_count += 1
            continue
            
        # Check duplicates
        if clean_email.lower() in seen_emails:
            skipped_duplicates += 1
            continue
        seen_emails.add(clean_email.lower())

        # Infer company if missing
        comp = company.strip() if company and company.strip() else _domain_to_company(clean_email)
        name = hr_name.strip() if hr_name and hr_name.strip() else None

        # Verify existence if enabled
        status = "VALID"
        reason = "Syntax valid"
        if verify_existence:
            v_res = verify_email_existence(comp, clean_email)
            status = v_res.status
            reason = v_res.reason

        if status == "INVALID":
            invalid_count += 1
        else:
            valid_count += 1
            processed_contacts.append({
                "company": comp,
                "hr_name": name,
                "email": clean_email,
                "status": status,
                "reason": reason
            })

    stats = {
        "total_found": total_emails_found,
        "valid_count": valid_count,
        "invalid_count": invalid_count,
        "duplicates_skipped": skipped_duplicates,
    }

    return processed_contacts, stats
